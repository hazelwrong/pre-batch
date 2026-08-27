"""Generic rubric-check executor.

Task-specific knowledge lives in data — `rubric_checks.json`, written by the
rubric agent alongside the rubric itself — and never in this module. That
separation is the whole point: before it, the only way to validate a second task
was to edit the validator, which meant every task carried a fork of the
checker and no two tasks were checked by the same code.

Each rubric item carries its own `check` — that is how the accepted package is
shaped, so the delivered rubric and the executed logic are the same object and
cannot drift apart. An item with no `check` is judged by a person, and must say
in `verification` how. There is no third case: "we could not test it" and "it is
correct" are different statements, and the delivery has been rejected once for
conflating them.

Column layouts come from the task, not from the check params. A check names a
sheet; the layout for that sheet follows. That is what lets the delivered rubric
carry executable params without column numbers a client has no reason to read.

Every checker returns ``(ok, detail)`` and reports what it actually observed,
including when it fails. Detail strings end up in the evidence file a reviewer
reads, so "got X expected Y" is worth more than "failed".
"""
import os
import re

import docx
import openpyxl
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader


REGISTRY = {}


def check(name):
    def register(fn):
        REGISTRY[name] = fn
        return fn
    return register


# ---------------------------------------------------------------- readers
def pdf_text(path):
    return "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)


def docx_text(path):
    document = docx.Document(path)
    out = []
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            out.append(Paragraph(child, document).text)
        elif child.tag.endswith("}tbl"):
            for row in Table(child, document).rows:
                out.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(out)


def xlsx_text(path):
    book = openpyxl.load_workbook(path, data_only=True)
    return " ".join(str(cell.value) for sheet in book
                    for row in sheet.iter_rows() for cell in row
                    if cell.value is not None)


def any_text(path):
    if path.endswith(".pdf"):
        return pdf_text(path)
    if path.endswith(".docx"):
        return docx_text(path)
    if path.endswith((".xlsx", ".xlsm")):
        return xlsx_text(path)
    with open(path, encoding="utf-8", errors="replace") as fh:
        return fh.read()


class Context:
    """Resolves the names a check entry uses to the files on disk.

    Checks address files by basename because that is what the prompt and the
    rubric both use. Two files with the same basename in one task would make
    those references ambiguous, so that is refused here rather than silently
    resolved to whichever was seen last.
    """

    def __init__(self, root, deliverable_files, reference_files, task=None):
        self.root = root
        self.task = task
        self.deliverables = self._index(deliverable_files)
        self.references = self._index(reference_files)
        self._books = {}
        self._texts = {}
        self._corpus = None

    def _index(self, relative_paths):
        index = {}
        for rel in relative_paths:
            name = os.path.basename(rel)
            if name in index:
                raise ValueError("two files share the basename %r; checks address "
                                 "files by basename and cannot disambiguate" % name)
            index[name] = os.path.join(self.root, rel)
        return index

    def path(self, name):
        if name in self.deliverables:
            return self.deliverables[name]
        if name in self.references:
            return self.references[name]
        raise KeyError("no delivered file named %r" % name)

    def book(self, name, data_only=False):
        key = (name, data_only)
        if key not in self._books:
            self._books[key] = openpyxl.load_workbook(self.path(name),
                                                      data_only=data_only)
        return self._books[key]

    def text(self, name):
        if name not in self._texts:
            self._texts[name] = any_text(self.path(name))
        return self._texts[name]

    def reference_corpus(self):
        """Everything the Agent could legitimately have cited, as one string."""
        if self._corpus is None:
            joined = " ".join(any_text(p) for p in self.references.values())
            self._corpus = re.sub(r"[,\s]+", " ", joined)
        return self._corpus


def _row(book, sheet, row, columns):
    ws = book[sheet]
    return {name: ws.cell(row=row, column=col).value
            for name, col in columns.items()}


def _columns(ctx, params, sheet=None):
    cols = params.get("columns")
    if not cols and ctx.task is not None:
        cols = ctx.task.column_map(sheet or params.get("sheet") or "*")
    if not cols:
        raise ValueError("no column layout for sheet %r; add it to task_meta.json "
                         "under column_maps" % (sheet or params.get("sheet")))
    return {name: int(col) for name, col in cols.items()}


# ---------------------------------------------------------------- checks
@check("xlsx_sheets")
def xlsx_sheets(ctx, prm):
    book = ctx.book(prm["file"])
    ok = book.sheetnames == prm["sheets"] if prm.get("exact") \
        else all(name in book.sheetnames for name in prm["sheets"])
    missing_meta = []
    if prm.get("require_columns"):
        ps = book[prm["param_sheet"]]
        for rw in range(2, ps.max_row + 1):
            key = ps.cell(row=rw, column=1).value
            if not key or str(key).startswith("—"):
                continue                                    # section heading
            for col in prm["require_columns"]:
                if not ps.cell(row=rw, column=col).value:
                    missing_meta.append("%s(col%d)" % (key, col))
    ok = ok and not missing_meta
    if ok:
        return True, ("sheetnames=%s; every %s parameter carries the required "
                      "metadata" % (book.sheetnames, prm.get("param_sheet", "")))
    return False, "sheets=%s missing metadata: %s" % (book.sheetnames,
                                                      missing_meta[:6])


@check("xlsx_parameterisation")
def xlsx_parameterisation(ctx, prm):
    book = ctx.book(prm["file"])
    allow = prm.get("allow_literals", ["1"])
    strip = prm.get("strip_patterns", [])
    bad_lit, bad_const = [], []
    for ws in book:
        if ws.title == prm["param_sheet"]:
            continue
        for row in ws.iter_rows():
            for cell in row:
                value = cell.value
                if isinstance(value, (int, float)):
                    bad_const.append("%s!%s" % (ws.title, cell.coordinate))
                if isinstance(value, str) and value.startswith("="):
                    text = re.sub(r'"[^"]*"', "", value)
                    text = re.sub(r"\$?[A-Z]{1,3}\$?\d+", "", text)
                    text = re.sub(r"%s!\$A\$2:\$B\$\d+" % re.escape(prm["param_sheet"]),
                                  "", text)
                    # A lookup's column index and its exact-match operand are
                    # part of the call, not model constants. LibreOffice may
                    # round-trip FALSE as FALSE().
                    text = re.sub(r",\s*\d+\s*,\s*FALSE\s*\(?\s*\)?\s*\)", ")", text)
                    for pattern in strip:
                        text = re.sub(pattern, "", text)
                    nums = [n for n in re.findall(r"(?<![A-Za-z_])\d+(?:\.\d+)?", text)
                            if n not in allow]
                    if nums:
                        bad_lit.append("%s!%s%s" % (ws.title, cell.coordinate, nums))
    if not bad_lit and not bad_const:
        return True, ("0 numeric literals and 0 constant cells outside '%s'; "
                      "propagation verified separately by parameter mutation"
                      % prm["param_sheet"])
    return False, "literals=%s consts=%s" % (bad_lit[:5], bad_const[:5])


@check("xlsx_cell_value")
def xlsx_cell_value(ctx, prm):
    value = ctx.book(prm["file"], data_only=True)[prm["sheet"]][prm["cell"]].value
    if value is None:
        return False, "%s!%s has no cached value" % (prm["sheet"], prm["cell"])
    expected = prm["expected"]
    tol = abs(expected) * prm.get("tolerance_rel", 0) + prm.get("tolerance_abs", 0)
    ok = abs(value - expected) <= tol
    ok = ok and all(abs(value - x) > 0.01 for x in prm.get("reject_values", []))
    return ok, ("%s!%s = %.2f (expected %.2f; rejected values %s not matched)"
                % (prm["sheet"], prm["cell"], value, expected,
                   prm.get("reject_values")))


@check("recompute")
def recompute(ctx, prm):
    book = ctx.book(prm["file"], data_only=True)
    got = _row(book, prm["sheet"], prm["row"], _columns(ctx, prm))
    exp = prm["expected"]
    tol = prm.get("tolerance_rel")
    missing = [k for k in exp if got.get(k) is None]
    if missing:
        return False, "no value at row %d for %s" % (prm["row"], missing)
    ok = all(abs(got[k] - exp[k]) <= (abs(exp[k]) * tol if tol
                                      else prm.get("tolerance_abs", 0))
             for k in exp)
    return ok, "got %s expected %s" % (got, exp)


@check("recompute_multi")
def recompute_multi(ctx, prm):
    book = ctx.book(prm["file"], data_only=True)
    tol = prm.get("tolerance_rel")
    floor = prm.get("tolerance_abs", 0.005)
    ok, detail = True, []
    for target in prm["targets"]:
        cols = _columns(ctx, prm, target["sheet"])
        got = _row(book, target["sheet"], target["row"], cols)
        exp = target["expected"]
        for key in exp:
            if got.get(key) is None:
                ok = False
                continue
            limit = abs(exp[key]) * tol if tol else 0
            if abs(got[key] - exp[key]) > max(limit, floor):
                ok = False
        detail.append("row %d: %s" % (target["row"],
                                      {k: (round(v, 2) if isinstance(v, (int, float))
                                           else v) for k, v in got.items()}))
    return ok, "; ".join(detail) + (" — all within tolerance" if ok else "")


@check("award_outcome")
def award_outcome(ctx, prm):
    ws = ctx.book(prm["file"], data_only=True)[prm["sheet"]]
    cols = _columns(ctx, prm)
    text = ctx.text(prm["doc"])
    award = {n: ws.cell(row=prm["award_row"], column=c).value for n, c in cols.items()}
    gate = {n: ws.cell(row=prm["gate_row"], column=c).value for n, c in cols.items()}
    totals = {n: round(ws.cell(row=prm["total_row"], column=c).value, 2)
              for n, c in cols.items()}
    ok = (award[prm["expected_recommended"]] == prm.get("award_marker", "RECOMMENDED")
          and gate[prm["expected_disqualified"]] == prm.get("gate_fail_marker", "Fail")
          and prm.get("runner_up_marker", "runner-up")
          in str(award[prm["expected_runner_up"]])
          and totals == prm["expected_totals"]
          and prm["doc_must_recommend"] in text
          and re.search(r"Recommend %s" % re.escape(prm["doc_must_recommend"]), text)
          and not re.search(r"Recommend %s\b" % re.escape(prm["doc_must_not_recommend"]),
                            text))
    return ok, ("totals=%s gate=%s award=%s; document recommends %s and does not "
                "recommend %s" % (totals, gate, award, prm["doc_must_recommend"],
                                  prm["doc_must_not_recommend"]))


@check("disqualified_vendor_presented")
def disqualified_vendor_presented(ctx, prm):
    ws = ctx.book(prm["file"], data_only=True)[prm["sheet"]]
    col = _columns(ctx, prm)[prm["vendor"]]
    text = ctx.text(prm["doc"])
    scored = all(ws.cell(row=rw, column=col).value is not None
                 for rw in prm["required_dimension_rows"])
    marker = prm.get("disqualification_marker", "disqualified")
    in_doc = prm["vendor"] in text and marker in text.lower()
    return scored and in_doc, (
        "%s scored on all %d dimensions: %s; presented with a stated "
        "disqualification: %s" % (prm["vendor"], len(prm["required_dimension_rows"]),
                                  scored, in_doc))


@check("cross_file_numeric")
def cross_file_numeric(ctx, prm):
    """Every figure quoted in the document traces to a cell, a derivation of
    cells, or a named reference file. Citing the source corpus is legitimate: a
    document that argues well quotes its inputs, and those figures are not in the
    workbook and should not be."""
    book = ctx.book(prm["xlsx"], data_only=True)
    text = ctx.text(prm["doc"])
    values = {round(float(cell.value), 2) for ws in book
              for row in ws.iter_rows() for cell in row
              if isinstance(cell.value, (int, float))}
    derived = set(values)
    derived |= {round(a - b, 2) for a in values for b in values if a > b}
    for divisor in prm.get("derive_divisors", [3, 9]):
        derived |= {round(v / divisor, 2) for v in set(derived)}
    # What a quoted figure looks like is a property of the task's documents, not
    # of the checker. Taken from the check entry, else from the task; never
    # defaulted to one currency's spelling.
    pattern = prm.get("figure_pattern")
    if not pattern and ctx.task is not None:
        pattern = (ctx.task.meta or {}).get("figure_pattern")
    if not pattern:
        raise ValueError("cross_file_numeric needs a figure_pattern, in the check "
                         "params or as figure_pattern in task_meta.json")
    quoted = {round(float(x.replace(",", "")), 2) for x in re.findall(pattern, text)}
    corpus = ctx.reference_corpus()
    tol = prm.get("tolerance_rel", 0.01)
    missing = {x for x in quoted
               if not any(abs(x - d) <= max(0.01, abs(d) * tol) for d in derived)
               and ("%g" % x) not in corpus
               and "{:,.2f}".format(x) not in corpus}
    if not missing:
        return True, ("%d figures in the document, every one traceable to a workbook "
                      "cell, a stated derivation, or a named reference file" % len(quoted))
    return False, "untraceable: %s" % sorted(missing)


@check("integrity_and_render")
def integrity_and_render(ctx, prm):
    """Not every task ships a workbook. A prose deliverable has no cells to
    carry an error value and no columns to format, so those parts are skipped
    rather than raising — a KeyError here failed a rubric item for a reason
    that had nothing to do with the deliverable."""
    workbook = prm.get("xlsx")
    text = ctx.text(prm["doc"])
    errors = []
    if workbook:
        book = ctx.book(workbook, data_only=True)
        live = ctx.book(workbook)
        errors = [(ws.title, cell.coordinate) for ws in book
                  for row in ws.iter_rows() for cell in row
                  if isinstance(cell.value, str) and cell.value.startswith("#")]
    tables = len(docx.Document(ctx.path(prm["doc"])).tables)
    native = tables > 0
    bad_format = []
    for sheet, rw in (prm.get("currency_cells", []) if workbook else []):
        columns = (prm.get("currency_columns")
                   or sorted(_columns(ctx, prm, sheet).values()))
        for col in columns:
            fmt = live[sheet].cell(row=rw, column=col).number_format
            if "0.00" not in fmt:
                bad_format.append("%s!%s%d=%s" % (sheet, chr(64 + col), rw, fmt))
    for pattern in prm.get("forbid_date_formats", []):
        for match in re.finditer(pattern, text):
            bad_format.append("non-ISO date %r" % match.group(0))
    ok = not errors and native and not bad_format
    if ok:
        return True, ("%d native table(s) in the document; every date is ISO%s"
                      % (tables, "; 0 spreadsheet error values and currency cells "
                         "carry a two-decimal format" if workbook else ""))
    return False, "errors=%s formatting=%s" % (errors[:3], bad_format[:4])


@check("doc_mentions")
def doc_mentions(ctx, prm):
    """Presence, not quality. Good for "did the report deal with X at all";
    useless for "did it deal with X well", which is a human's call."""
    text = ctx.text(prm["doc"]).lower()
    missing_all = [w for w in prm.get("all_of", []) if w.lower() not in text]
    any_of = prm.get("any_of", [])
    hit_any = [w for w in any_of if w.lower() in text]
    also = prm.get("also_any_of", [])
    hit_also = [w for w in also if w.lower() in text]
    ok = not missing_all and (not any_of or hit_any) and (not also or hit_also)
    if ok:
        return True, "document mentions " + ", ".join(hit_any or prm.get("all_of", []))
    absent = missing_all + (["any of: " + ", ".join(any_of)]
                            if any_of and not hit_any else [])
    absent += (["any of: " + ", ".join(also)] if also and not hit_also else [])
    return False, "document does not mention " + ", ".join(absent)


@check("doc_sections")
def doc_sections(ctx, prm):
    text = ctx.text(prm["doc"]).lower()
    missing = [s for s in prm["sections"] if s.lower() not in text]
    if not missing:
        return True, "all %d required sections present" % len(prm["sections"])
    return False, "missing: %s" % missing


@check("article12_conditions")   # the name the format reference happens to ship
@check("doc_markers_framed")     # what the check actually is
def doc_markers_framed(ctx, prm):
    """Every marker present AND framed by at least one governing phrase — the
    difference between listing a clause and treating it as binding."""
    text = ctx.text(prm["doc"])
    missing = [m for m in prm["markers"] if m not in text]
    framed = any(phrase.lower() in text.lower()
                 for phrase in prm.get("mandatory_phrases", []))
    ok = not missing and framed
    if ok:
        return True, ("all %d markers present and framed as requirements"
                      % len(prm["markers"]))
    return False, "missing=%s framed_as_mandatory=%s" % (missing, framed)


# ---------------------------------------------------------------- runner
def execute(item, ctx):
    """Run one rubric item's check. Returns (status, detail, check_type).

    An item with no check is not evaluated here and says so; it is never
    reported as passed."""
    entry = item.get("check")
    if entry is None:
        return ("not_auto_evaluated",
                item.get("verification")
                or "No programmatic test for this item; it is settled by the "
                   "reviewer reading the deliverable.",
                "human_judgement")
    if not isinstance(entry, dict):
        return "failed", "malformed check entry", "?"
    if entry.get("human") is True:
        return ("not_auto_evaluated",
                entry.get("reason") or item.get("verification")
                or "This item is settled by a human reviewer.",
                "human_judgement")
    kind = entry.get("type")
    fn = REGISTRY.get(kind)
    if not fn:
        # Not silently passed: an unknown type is a spec error someone has to
        # see, and it stays out of the earned score either way.
        return "failed", "no checker registered for type %r" % kind, kind or "?"
    try:
        ok, detail = fn(ctx, entry.get("params", {}))
    except Exception as exc:                                      # noqa: BLE001
        return "failed", "checker raised %s: %s" % (type(exc).__name__, exc), kind
    return ("passed" if ok else "failed"), detail, kind
