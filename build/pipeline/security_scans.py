"""The five scans §9 asks for beyond answer leakage.

Requirement §9 lists 答案泄漏、隐私、版权、恶意文件、路径穿越和密钥扫描 — six
scan types. Only the first was being produced; the others existed as assertions
in the provenance manifest ("contains_pii": false) rather than as artefacts.
An assertion is not a scan: nothing was actually looked at.

Each function here returns a result dict that is written to validation evidence,
so a reviewer can see what was searched for and what was found.
"""
import hashlib
import os
import re
import zipfile

import docx
import openpyxl

try:
    from pypdf import PdfReader
except ImportError:  # PDF support is required only when a PDF is actually scanned.
    PdfReader = None


def _suffix(path):
    return os.path.splitext(os.fspath(path))[1].lower()


def _text_of(path):
    path = os.fspath(path)
    suffix = _suffix(path)
    if suffix == ".pdf":
        if PdfReader is None:
            raise RuntimeError("pypdf is required to scan PDF files")
        return "\n".join(p.extract_text() or "" for p in PdfReader(path).pages)
    if suffix == ".docx":
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        parts += [c.text for t in d.tables for r in t.rows for c in r.cells]
        return "\n".join(parts)
    if suffix in (".xlsx", ".xlsm"):
        # Formula text is security-relevant. data_only=True exposes only cached
        # results, which made WEBSERVICE/DDE formulas invisible to the scanner.
        wb = openpyxl.load_workbook(path, data_only=False)
        return "\n".join(str(c.value) for ws in wb for row in ws.iter_rows()
                         for c in row if c.value is not None)
    try:
        return open(path, encoding="utf-8", errors="replace").read()
    except OSError:
        return ""


def _cached_text(path, text_cache):
    key = os.fspath(path)
    if text_cache is None:
        return _text_of(key)
    if key not in text_cache:
        text_cache[key] = _text_of(key)
    return text_cache[key]


# --------------------------------------------------------------------------
# 1. Privacy / PII
# --------------------------------------------------------------------------
PII = [
    ("mainland ID number", r"\b[1-9]\d{5}(19|20)\d{2}(0[1-9]|1[0-2])"
                           r"(0[1-9]|[12]\d|3[01])\d{3}[\dXx]\b"),
    ("mainland mobile", r"(?<!\d)1[3-9]\d{9}(?!\d)"),
    ("email address", r"\b[\w.+-]+@[\w-]+\.[\w.]{2,}\b"),
    ("bank card", r"(?<!\d)(?:\d[ -]?){15,18}\d(?!\d)"),
    ("passport", r"\b[EGDSPH]\d{8}\b"),
    ("street address", r"\b\d{1,4}\s+[A-Z][a-z]+\s+(Street|Road|Avenue|Lane)\b"),
    # The list above was mainland-only, so a UK council timetable carrying
    # seven personal mobile numbers scanned clean. A privacy scanner that
    # cannot see the numbers in the material it is pointed at reports the
    # opposite of the truth, which is worse than not running it.
    ("UK mobile", r"(?<!\d)07\d{3}\s?\d{3}\s?\d{3}(?!\d)"),
    ("UK landline", r"(?<!\d)0(?:1\d{3}|2\d)\s{1,2}\d{3}\s?\d{3}(?!\d)"),
    ("intl E.164", r"\+\d{1,3}[\s-]?\d{2,4}[\s-]?\d{3,4}[\s-]?\d{3,4}"),
    ("NANP phone", r"(?<!\d)(?:\(\d{3}\)\s?|\d{3}[-.])\d{3}[-.]\d{4}(?!\d)"),
]


def scan_privacy(files, note=None, text_cache=None):
    hits = []
    for f in files:
        t = _cached_text(f, text_cache)
        for label, pat in PII:
            for m in re.finditer(pat, t):
                hits.append({"file": os.path.basename(f), "type": label,
                             "match": m.group(0)[:6] + "…"})
    return {"passed": not hits, "patterns_checked": len(PII),
            "files_scanned": len(files), "hits": hits,
            # Whose names these are is a fact about the task's material, not
            # about the scanner. Hardcoding it left a retail task's sentence
            # ("store managers") sitting in a federal regulatory package whose
            # names are real public officials — the note said the opposite of
            # the truth, and a reviewer caught it.
            # With nothing found there is nothing to interpret; a standing
            # sentence there reads as an excuse prepared in advance.
            "note": (note or "No interpretation of the findings is recorded; the "
                             "task must account for each hit.") if hits else None}


# --------------------------------------------------------------------------
# 2. Copyright / third-party material
# --------------------------------------------------------------------------
COPYRIGHT = [
    ("copyright notice", r"©|\(c\)\s*\d{4}|Copyright\s+\d{4}"),
    ("all rights reserved", r"[Aa]ll [Rr]ights [Rr]eserved"),
    ("licence marker", r"\b(GPL|LGPL|MIT License|Apache License|CC BY|"
                       r"Creative Commons|proprietary and confidential)\b"),
    ("trademark", r"[™®]"),
]


def scan_copyright(files, text_cache=None):
    hits = []
    for f in files:
        t = _cached_text(f, text_cache)
        for label, pat in COPYRIGHT:
            for m in re.finditer(pat, t):
                hits.append({"file": os.path.basename(f), "type": label,
                             "match": m.group(0)})
    return {"passed": not hits, "patterns_checked": len(COPYRIGHT),
            "files_scanned": len(files), "hits": hits,
            "note": ("This pattern scan only reports embedded rights notices. "
                     "Ownership, licence and source type are evaluated from the "
                     "task-specific provenance records.")}


# --------------------------------------------------------------------------
# 3. Malicious content
# --------------------------------------------------------------------------
def scan_malicious(files, text_cache=None):
    hits = []
    for f in files:
        suffix = _suffix(f)
        if suffix in (".docx", ".xlsx", ".xlsm", ".pptx"):
            with zipfile.ZipFile(f) as z:
                names = z.namelist()
            for n in names:
                if n.startswith("word/vbaProject") or n.endswith(".bin") \
                        or "vbaProject" in n or n.startswith("xl/macroSheets"):
                    hits.append({"file": os.path.basename(f),
                                 "type": "macro or binary part", "part": n})
                if n.startswith("word/embeddings") or n.startswith("xl/embeddings"):
                    hits.append({"file": os.path.basename(f),
                                 "type": "embedded object", "part": n})
            t = _cached_text(f, text_cache)
            for pat, label in [(r"=\s*(cmd|powershell|WEBSERVICE|DDE)\b",
                                "external-call formula"),
                               (r"\bAuto_Open\b|\bWorkbook_Open\b", "auto-run macro")]:
                for m in re.finditer(pat, t, re.I):
                    hits.append({"file": os.path.basename(f), "type": label,
                                 "match": m.group(0)})
        if suffix == ".pdf":
            raw = open(f, "rb").read()
            for marker, label in [(b"/JavaScript", "embedded JavaScript"),
                                  (b"/JS", "embedded JavaScript"),
                                  (b"/Launch", "launch action"),
                                  (b"/EmbeddedFile", "embedded file"),
                                  (b"/OpenAction", "open action")]:
                if marker in raw:
                    hits.append({"file": os.path.basename(f), "type": label})
    return {"passed": not hits, "files_scanned": len(files), "hits": hits,
            "note": "Office packages are inspected part by part for macros, "
                    "embedded objects and external-call formulas; PDFs for "
                    "JavaScript, launch actions and embedded files."}


# --------------------------------------------------------------------------
# 4. Path traversal / absolute paths
# --------------------------------------------------------------------------
PATHS = [
    ("parent traversal", r"\.\.[\\/]"),
    ("POSIX absolute path", r"(?<![\w.])/(?:Users|home|tmp|var|etc|mnt|opt)/[\w./-]+"),
    ("Windows absolute path", r"[A-Za-z]:\\[\\\w.-]+"),
    ("file URI", r"file://"),
    ("UNC path", r"\\\\[\w.-]+\\"),
]


def scan_paths(files, declared_paths, text_cache=None):
    hits = []
    for f in files:
        t = _cached_text(f, text_cache)
        for label, pat in PATHS:
            for m in re.finditer(pat, t):
                hits.append({"file": os.path.basename(f), "type": label,
                             "match": m.group(0)[:80]})
    bad_declared = [p for p in declared_paths
                    if p.startswith("/") or ".." in p or "\\" in p
                    or re.match(r"^[A-Za-z]:", p)]
    return {"passed": not hits and not bad_declared,
            "patterns_checked": len(PATHS), "files_scanned": len(files),
            "content_hits": hits, "declared_path_violations": bad_declared}


# --------------------------------------------------------------------------
# 5. Secrets / credentials
# --------------------------------------------------------------------------
SECRETS = [
    ("AWS access key", r"\bAKIA[0-9A-Z]{16}\b"),
    ("generic API key", r"\b(?:api[_-]?key|apikey|secret|token|passwd|password)"
                        r"\s*[:=]\s*['\"]?[A-Za-z0-9/+_-]{12,}"),
    ("bearer token", r"\bBearer\s+[A-Za-z0-9._-]{20,}"),
    ("private key block", r"-----BEGIN [A-Z ]*PRIVATE KEY-----"),
    ("connection string", r"\b(?:mongodb|postgres|postgresql|mysql|redis)://[^\s]+"),
    ("JWT", r"\bey[A-Za-z0-9_-]{10,}\.[A-Za-z0-9_-]{10,}\.[A-Za-z0-9_-]{10,}"),
]


def scan_secrets(files, text_cache=None):
    hits = []
    for f in files:
        t = _cached_text(f, text_cache)
        for label, pat in SECRETS:
            for m in re.finditer(pat, t, re.I):
                hits.append({"file": os.path.basename(f), "type": label,
                             "match": m.group(0)[:12] + "…"})
    return {"passed": not hits, "patterns_checked": len(SECRETS),
            "files_scanned": len(files), "hits": hits}


def scanned_subjects(files):
    """Record what was scanned, by content, not just by name.

    A scan result that lists only filenames and counts looks identical whether
    it was re-run this morning or copied from the first build months ago. That
    is exactly what happened here: the evidence was regenerated on every build,
    but nothing in it could change, so its hash never moved and a reviewer had
    no way to tell. Binding each result to the SHA-256 of the bytes it examined
    makes staleness visible: change a reference file and this section changes
    with it, or the evidence is provably out of date.
    """
    out = []
    for f in sorted(files):
        if not os.path.isfile(f):
            continue
        with open(f, "rb") as fh:
            h = hashlib.sha256(fh.read()).hexdigest()
        out.append({"file": os.path.basename(f), "sha256": h,
                    "bytes": os.path.getsize(f)})
    return out


def run_all(files, declared_paths, extra_subjects=None, notes=None):
    subjects = scanned_subjects(files)
    if extra_subjects:
        subjects = subjects + list(extra_subjects)
    digest = hashlib.sha256(
        "".join(x["sha256"] for x in subjects).encode()).hexdigest()
    text_cache = {}
    return {
        "scan_subjects": subjects,
        "scan_subject_digest": digest,
        "scan_subject_digest_note": (
            "SHA-256 over the hashes of everything scanned, in name order. If any "
            "scanned file changes, this digest changes, so evidence produced "
            "before a change cannot silently pass as evidence produced after it."),
        "privacy_pii": scan_privacy(
            files, (notes or {}).get("privacy_pii"), text_cache),
        "copyright": scan_copyright(files, text_cache),
        "malicious_content": scan_malicious(files, text_cache),
        "path_traversal": scan_paths(files, declared_paths, text_cache),
        "secrets": scan_secrets(files, text_cache),
    }
